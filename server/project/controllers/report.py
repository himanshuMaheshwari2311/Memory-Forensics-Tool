from project import app
from flask import render_template, request, session, send_from_directory
from flask_wtf import FlaskForm
from docx import Document
from docx.shared import Inches
import comtypes.client
import os
import json
import time

@app.route('/update_report', methods = ['GET', 'POST'])
def update_report():
	print "update_report"
	if request.method == 'POST':
		print "POST"
		overview = request.form['case_overview']
		acquisition = request.form['case_acquisition']
		findings = request.form['case_findings']
		conclusion = request.form['case_conclusion']
		print "after extraction"
		resp = {}
		case_data = {}
		print "before with 1"
		with open('../data/json/' + session['selected_case'], mode='r') as f:
			case_data = json.load(f)
			case_data['case_overview'] = overview
			case_data['case_acquisition'] = acquisition
			case_data['case_findings'] = findings
			case_data['case_conclusion'] = conclusion

		print "before with 2"
		with open('../data/json/' + session['selected_case'], mode='w') as f:
			json.dump(case_data, f, indent = 4)

		resp['result'] = True
		print "before return"
		return json.dumps(resp)

@app.route('/get_report', methods=['GET', 'POST'])
def get_report():
	localtime = time.asctime( time.localtime(time.time()))
	report_name = session['selected_case'].split('.')[0] + '_report.docx'
	pdf_report_name = session['selected_case'].split('.')[0] + '_report.pdf'
	document = Document()
	document.add_heading(session['selected_case'].split('.')[0] + " Report", 0)
	lines = ""	
	with open('../data/json/' + session['selected_case'], mode='r') as f:
		case_data = json.load(f)
		document.add_heading('Case Overview', level=1)
		document.add_paragraph('\t' + case_data['case_overview'])
		document.add_heading('Case Acquisition', level=1)
		document.add_paragraph('\t' + case_data['case_acquisition'])
		document.add_heading('Case Findings', level=1)
		document.add_paragraph('\t' + case_data['case_findings'])
		document.add_heading('Case Conclusion', level=1)
		document.add_paragraph('\t' + case_data['case_conclusion'])
		document.add_heading('Report Generation Time', level=1)
		document.add_paragraph('\t' + str(localtime))
		document.add_heading('Artifacts', level=1)
		i = 0
		hr = "---------------------------------------------------------------------------------------------"
		for artifact in case_data['artifacts']:
			key = next(iter(artifact))
			document.add_heading(key, level=2)
			document.add_paragraph(hr)
			value = artifact[key]
			j = 0
			for module in value:
				if module['marked'] == "disabled":
					for ke, va in module.iteritems():
						lines += str(ke) + ": " + str(va) + "\n"
						if ke == 'comment':
							lines += '\n'
					lines += hr
					document.add_paragraph(lines)
					lines = ''
				j += 1
			i += 1
	p = document.add_paragraph(lines)
	document.add_page_break()
	document.save('../data/pdfs/' + report_name)
	file_path = os.getcwd() + '\\..\\data\\pdfs\\'
	word = comtypes.client.CreateObject('Word.Application')
	doc = word.Documents.Open(file_path + report_name)
	doc.SaveAs(file_path + pdf_report_name, FileFormat=17)
	doc.Close()
	word.Quit()
	try:
		return send_from_directory('..\\..\\data\\pdfs', pdf_report_name)
	except Exception as e:
		return str(e)