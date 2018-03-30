from project import app
from flask import render_template, request, session, send_from_directory
from flask_wtf import FlaskForm
from docx import Document
from docx.shared import Inches
import comtypes.client
import  os

import json

@app.route('/case', methods = ['GET', 'POST'])
def case():
	print "Case!!!"
	resp = {}
	if 'username' not in session:
		resp['result_type'] = "info"
		resp['result'] = "Please log in!"
		return render_template('accounts/login.html', resp = resp)
	if request.method == 'GET':
		case_name = request.args.get('case_name')
		if case_name is None:
			resp['result_type'] = "danger"
			resp['result'] = "Please specify the case!"
			resp['cases'] = session['cases']
			return render_template('accounts/index.html', resp = resp)
		try:
			if case_name not in session['cases']:
				print "Case name not in session"
				resp['result_type'] = "danger"
				resp['result'] = "Specified case name is wrong!"
				resp['cases'] = session['cases']
				return render_template('accounts/index.html', resp = resp)
			
			# added here
			try:
				case_data = json.load(open('../data/json/' + case_name))
			except Exception as e:
				print "Could not open file"
				return render_template('accounts/index.html', resp = resp)
			session['selected_case'] = case_name	
			return render_template('case/case.html', case_data = case_data)
		
		except Exception as e:
			print "Exception: " + str(e)
			resp['result_type'] = "danger"
			resp['result'] = "Specified case name is wrong!"
			resp['cases'] = session['cases']
			return render_template('accounts/index.html', resp = resp)

# add function
@app.route('/add_artifact', methods = ['GET', 'POST'])
def add_artifact():
    if request.method == 'POST':
		object_id = request.form['object_id']
		resp = {}
		with open('../data/json/' + session['selected_case'], mode='r') as f:
			case_data = json.load(f)
			new_case_data = case_data
			i = 0
			for artifact in case_data['artifacts']:
				key = next(iter(artifact))
				value = artifact[key]
				j = 0
				for module in value:
					if str(module['object_id']) == str(object_id) and module['marked'] == "enabled":
						new_case_data['artifacts'][i][key][j]['marked'] = "disabled"
						if 'comment' in request.form:
							new_case_data['artifacts'][i][key][j]['comment'] = request.form['comment']
						resp['obj'] = new_case_data['artifacts'][i][key][j]
					j += 1
				i += 1
		with open('../data/json/' + session['selected_case'], mode='w') as f:
			json.dump(new_case_data, f, indent = 4)

		resp['result'] = True
		return json.dumps(resp)

# remove function
@app.route('/remove_artifact', methods = ['GET', 'POST'])
def remove_artifact():
    if request.method == 'POST':
		object_id = request.form['object_id']
		resp = {}
		with open('../data/json/' + session['selected_case'], mode='r') as f:
			case_data = json.load(f)
			new_case_data = case_data
			i = 0
			for artifact in case_data['artifacts']:
				key = next(iter(artifact))
				value = artifact[key]
				j = 0
				for module in value:
					if str(module['object_id']) == str(object_id) and module['marked'] == "disabled":
						new_case_data['artifacts'][i][key][j]['marked'] = "enabled"
						if 'comment' in new_case_data['artifacts'][i][key][j]:
							new_case_data['artifacts'][i][key][j]['comment'] = ""
					j += 1
				i += 1
		with open('../data/json/' + session['selected_case'], mode='w') as f:
			json.dump(new_case_data, f, indent = 4)

		resp['result'] = True
		return json.dumps(resp)

@app.route('/update_report', methods = ['GET', 'POST'])
def update_report():
	print "update_report"
	if request.method == 'POST':
		print "POST"
		overview = request.form['case_overview']
		acquisition = request.form['case_acquisition']
		findings = request.form['case_findings']
		conclusion = request.form['case_conclusion']
		print "after extraction"
		resp = {}
		case_data = {}
		print "before with 1"
		with open('../data/json/' + session['selected_case'], mode='r') as f:
			case_data = json.load(f)
			case_data['case_overview'] = overview
			case_data['case_acquisition'] = acquisition
			case_data['case_findings'] = findings
			case_data['case_conclusion'] = conclusion

		print "before with 2"
		with open('../data/json/' + session['selected_case'], mode='w') as f:
			json.dump(case_data, f, indent = 4)

		resp['result'] = True
		print "before return"
		return json.dumps(resp)

@app.route('/get_report', methods=['GET', 'POST'])
def get_report():
	report_name = session['selected_case'].split('.')[0] + '_report.docx'
	pdf_report_name = session['selected_case'].split('.')[0] + '_report.pdf'
	document = Document()
	document.add_heading(session['selected_case'].split('.')[0] + " Report", 0)
	lines = ""	
	with open('../data/json/' + session['selected_case'], mode='r') as f:
		case_data = json.load(f)
		lines += 'Case Overview\n' + case_data['case_overview'] + '\n\n'
		lines += 'Case Acquisition\n' + case_data['case_acquisition'] + '\n\n'
		lines += 'Case Findings\n' + case_data['case_findings'] + '\n\n'
		lines += 'Case Conclusion\n' + case_data['case_conclusion'] + '\n\n'
		lines += 'Artifacts\n'
		i = 0
		for artifact in case_data['artifacts']:
			key = next(iter(artifact))
			lines += key + ': \n'
			value = artifact[key]
			j = 0
			for module in value:
				if module['marked'] == "disabled":
					for ke, va in module.iteritems():
						lines += "\t" + str(ke) + ": " + str(va) + "\n"
					lines +="\n"
				j += 1
			i += 1
	p = document.add_paragraph(lines)
	document.add_page_break()
	document.save('../data/pdfs/' + report_name)
	file_path = os.getcwd() + '\\..\\data\\pdfs\\'
	print file_path
	word = comtypes.client.CreateObject('Word.Application')
	doc = word.Documents.Open(file_path + report_name)
	doc.SaveAs(file_path + pdf_report_name, FileFormat=17)
	doc.Close()
	word.Quit()
	try:
		return send_from_directory('..\\..\\data\\pdfs', pdf_report_name)
	except Exception as e:
		return str(e)