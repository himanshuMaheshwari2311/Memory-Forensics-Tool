from project import app
from flask import render_template, request, session, send_from_directory
from flask_wtf import FlaskForm

import json

@app.route('/send_report', methods = ['GET', 'POST'])
def send_report():
	try:
		return send_from_directory('..\\..\\data\\pdfs', 'cl.pdf')
	except Exception as e:
		return str(e)

@app.route('/update_report', methods = ['GET', 'POST'])
def update_report():
	if request.method == 'POST':
		overview = request.form['case_overview']
		acquisition = request.form['case_acquisition']
		findings = request.form['case_findings']
		conclusion = request.form['case_conclusion']
		
		resp = {}
		case_data = {}
		
		with open('../data/json/' + session['selected_case'], mode='r') as f:
			case_data = json.load(f)
			case_data['case_overview'] = overview
			case_data['case_acquisition'] = acquisition
			case_data['case_findings'] = findings
			case_data['case_conclusion'] = conclusion
		with open('../data/json/' + session['selected_case'], mode='w') as f:
			json.dump(case_data, f, indent = 4)

		resp['result'] = True
		return json.dumps(resp)

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

document.add_page_break()

document.save('demo.docx')
